# backend/main.py
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass
import shutil
import uuid
from typing import Dict, Any, List
import tempfile
import zipfile
import requests
import base64

# Imports for document processing (adapted from the original main.py)
from docxtpl import DocxTemplate
from openpyxl import load_workbook
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
import subprocess
from PyPDF2 import PdfMerger
from datetime import datetime
import time

try:
    import win32com.client as win32
    WIN32_AVAILABLE = True
except Exception:
    WIN32_AVAILABLE = False

import docx
import re

app = FastAPI(title="Fichas Automáticas API")

# Setup CORS to allow the React frontend to communicate with this API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For local development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RUTA_FORMATOS = os.path.join(BASE_DIR, "formatos")
# En Vercel, solo podemos escribir en /tmp
RUTA_SALIDA = os.path.join(tempfile.gettempdir(), "fichas_salida")
os.makedirs(RUTA_FORMATOS, exist_ok=True)
os.makedirs(RUTA_SALIDA, exist_ok=True)

# === AUXILIARY FUNCTIONS ===
def calcular_edad(fecha_nacimiento_str: str) -> str:
    try:
        if not fecha_nacimiento_str: return ""
        fecha_nac = datetime.strptime(fecha_nacimiento_str, "%d-%m-%Y")
        hoy = datetime.today()
        edad = hoy.year - fecha_nac.year - ((hoy.month, hoy.day) < (fecha_nac.month, fecha_nac.day))
        return str(edad)
    except Exception:
        return ""

def escribir_celda(ws, celda_ref: str, valor: str):
    celda = ws[celda_ref]
    for rango in ws.merged_cells.ranges:
        if celda.coordinate in rango:
            celda_superior = ws.cell(rango.min_row, rango.min_col)
            celda_superior.value = valor
            return
    celda.value = valor

# === EXTRACTION LOGIC ===
@app.post("/api/extract")
async def extract_data_from_docx(file: UploadFile = File(...)):
    # Create a temporary file to save the uploaded word document
    temp_dir = tempfile.mkdtemp()
    temp_file_path = os.path.join(temp_dir, file.filename)
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    datos = {}
    try:
        doc = docx.Document(temp_file_path)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el archivo Word: {str(e)}")
        
    texto_completo = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: texto_completo.append(t)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                text = c.text.strip(' \n\r\t')
                if text and text not in texto_completo:
                    texto_completo.append(text)
                    
    # === LÓGICA DE EXTRACCIÓN MEJORADA ===
    # 1. Búsqueda exacta en tablas (100% Precisión para Plantilla V2)
    for t in doc.tables:
        for r in t.rows:
            if len(r.cells) >= 2:
                lbl = r.cells[0].text.strip().upper()
                val = r.cells[1].text.strip()
                if val:
                    if "NOMBRES PROCESADO" in lbl:
                        partes = val.split()
                        if len(partes) >= 2:
                            datos['primer_nombre'] = partes[0]
                            datos['segundo_nombre'] = " ".join(partes[1:])
                        elif len(partes) == 1:
                            datos['primer_nombre'] = partes[0]
                    elif "APELLIDOS PROCESADO" in lbl:
                        partes = val.split()
                        if len(partes) >= 2:
                            datos['primer_apellido'] = partes[0]
                            datos['segundo_apellido'] = " ".join(partes[1:])
                        elif len(partes) == 1:
                            datos['primer_apellido'] = partes[0]
                    elif "TIPO DOCUMENTO" in lbl: datos['tipo_documento'] = val.lower()
                    elif "NUMERO DOCUMENTO" in lbl: datos['cedula'] = val.replace(".", "").replace(",", "")
                    elif "EXPEDIDA EN" in lbl: datos['expedida_en'] = val
                    elif "FECHA EXPEDICIÓN" in lbl: datos['fecha_expedicion'] = val
                    elif "FECHA NACIMIENTO" in lbl: datos['fecha_nacimiento'] = val
                    elif "NACIONALIDAD" in lbl: datos['nacionalidad'] = val
                    elif "SEXO" in lbl: datos['sexo'] = val.lower()
                    
                    elif "DIRECCIÓN DE RESIDENCIA" in lbl: datos['direccion'] = val
                    elif "DEPARTAMENTO" in lbl: datos['departamento'] = val
                    elif "MUNICIPIO" in lbl: datos['municipio'] = val
                    elif "BARRIO" in lbl: datos['barrio'] = val
                    elif "VEREDA" in lbl: datos['vereda'] = val
                    elif "CELULAR" in lbl: datos['celular'] = val
                    elif "TELÉFONO FIJO" in lbl: datos['telefono'] = val
                    elif "CORREO ELECTRÓNICO" in lbl: datos['correo'] = val
                    elif "NIVEL EDUCATIVO" in lbl: datos['nivel_educativo'] = val.lower()
                    elif "ESTADO CIVIL" in lbl: datos['estado_civil'] = val.lower()
                    elif "NÚMERO DE HIJOS" in lbl: datos['n_hijos'] = val
                    elif "ESTRATO" in lbl: datos['estrato'] = val
                    elif "VIVIENDA" in lbl: datos['vivienda'] = val.lower()
                    elif "SITUACIÓN LABORAL" in lbl: datos['situacion_laboral'] = val.lower()
                    elif "CARGO" in lbl: datos['cargo'] = val
                    elif "INGRESO MENSUAL" in lbl: datos['ingreso_mensual'] = val
                    
                    elif "RUS" in lbl: datos['rus'] = val
                    elif "RUC" in lbl: datos['ruc'] = val
                    elif "LUGAR DE RECEPCIÓN" in lbl: datos['lugar_recepcion'] = val
                    elif "FECHA DE RECEPCIÓN" in lbl: datos['fecha_recepcion'] = val
                    elif "CONDUCTA PUNIBLE" in lbl: datos['conducta_punible'] = val
                    elif "NÚMERO DE PROCESO" in lbl: datos['numero_proceso'] = val
                    elif "FECHA Y HORA DE CAPTURA" in lbl: datos['fecha_hora_captura'] = val
                    elif "FISCAL" in lbl: datos['fiscal'] = val
                    elif "JUZGADO" in lbl: datos['juez'] = val
                    elif "PRIVADO DE LA LIBERTAD" in lbl: 
                        datos['privado_libertad'] = val.upper() in ("SI", "SÍ", "TRUE", "X", "1")
                    elif "CENTRO DE RECLUSIÓN" in lbl: datos['centro_reclusion'] = val

    # Extraer Hechos (Plantilla V2 Tabla 3)
    try:
        if len(doc.tables) >= 4:
            hechos_val = doc.tables[3].rows[1].cells[0].text.strip()
            if hechos_val:
                datos['resumen_hechos'] = hechos_val.replace("PARA EL DIA (No modificar esta línea)", "").strip()
    except Exception:
        pass

    # 2. Expresiones Regulares de Respaldo (por si es el formato antiguo de texto plano)
    texto_plano = "\n".join(texto_completo)
    
    if 'celular' not in datos:
        m_cel = re.search(r'CEL(?:U?LAR|\.)?\s*[:\-]?\s*(\d{7,10})', texto_plano, re.IGNORECASE)
        if m_cel: datos['celular'] = m_cel.group(1)
            
    if 'estrato' not in datos:
        m_estrato = re.search(r'ESTRATO\s*(\d)', texto_plano, re.IGNORECASE)
        if m_estrato: datos['estrato'] = m_estrato.group(1)
            
    if 'cedula' not in datos:
        m_cc = re.search(r'CC\s*([\d\.\,]+)\s+(?:DE\s+)?([A-Z\s]+)', texto_plano)
        if m_cc: 
            datos['cedula'] = m_cc.group(1).replace(".", "").replace(",", "")
            datos['expedida_en'] = m_cc.group(2).strip()
            datos['tipo_documento'] = 'cc'
            
    if 'fecha_nacimiento' not in datos:
        m_fn = re.search(r'F\.N\s*(\d{2})\s*[-/]?(\d{2})\s*[-/]?(\d{4})', texto_plano)
        if m_fn: 
            datos['fecha_nacimiento'] = f"{m_fn.group(1)}-{m_fn.group(2)}-{m_fn.group(3)}"
            
    if 'n_hijos' not in datos:
        m_hijos = re.search(r'HIJOS\s*[:\-]?\s*(\d+)', texto_plano, re.IGNORECASE)
        if m_hijos: datos['n_hijos'] = m_hijos.group(1)
            
    if 'fiscal' not in datos:
        m_fiscal = re.search(r'(FISCAL[A-Z\s0-9]+(?:\(GARANTIAS\)|\(CONOCIMIENTO\))?)', texto_plano)
        if m_fiscal: datos['fiscal'] = m_fiscal.group(1).strip()
            
    if 'juez' not in datos:
        m_juez = re.search(r'(JUZGADO[A-Z\s0-9]+)', texto_plano)
        if m_juez: datos['juez'] = m_juez.group(1).strip()
            
    if 'numero_proceso' not in datos:
        m_proc = re.search(r'\b(\d{21})\b|\b(\d{21}-\d+)\b', texto_plano)
        if m_proc: 
            datos['numero_proceso'] = m_proc.group(1) or m_proc.group(2)
            
    if 'primer_nombre' not in datos:
        m_nombres = re.search(r'(?:NOMBRE PROCESADO\n|CEDULA PROCESADO\n)+([A-Z\s]+)\n', texto_plano)
        if not m_nombres:
            m_nombres = re.search(r'\b\d{21}.*?\t([A-Z\s]+)', texto_plano)
        
        if m_nombres:
            partes = m_nombres.group(1).strip().split()
            if len(partes) >= 4:
                datos['primer_nombre'] = partes[0]
                datos['segundo_nombre'] = partes[1]
                datos['primer_apellido'] = partes[2]
                datos['segundo_apellido'] = " ".join(partes[3:])
            elif len(partes) == 3:
                datos['primer_nombre'] = partes[0]
                datos['primer_apellido'] = partes[1]
                datos['segundo_apellido'] = partes[2]
            elif len(partes) == 2:
                datos['primer_nombre'] = partes[0]
                datos['primer_apellido'] = partes[1]
                
    if 'resumen_hechos' not in datos:
        m_hechos = re.search(r'LOS SIGUIENTES HECHOS:\s*(.*?)(?:\nPARA EL DIA|\nLE EXPLIQUE|$)', texto_plano, re.DOTALL | re.IGNORECASE)
        if m_hechos:
            hechos_texto = m_hechos.group(1).strip()
            if hechos_texto: datos['resumen_hechos'] = hechos_texto
        
    # Clean up temp file
    try:
        shutil.rmtree(temp_dir)
    except Exception:
        pass
        
    return {"extracted_data": datos}

# === DOCUMENT GENERATION LOGIC ===

def llenar_excel1(plantilla, salida, datos):
    wb = load_workbook(plantilla)
    ws = wb.active

    escribir_celda(ws, "K5", str(datos.get("rus", "")))
    escribir_celda(ws, "W5", str(datos.get("ruc", "")))
    escribir_celda(ws, "D8", str(datos.get("lugar_recepcion", "")))
    escribir_celda(ws, "R8", str(datos.get("fecha_recepcion", "")))
    escribir_celda(ws, "D38", str(datos.get("primer_apellido", "")))
    escribir_celda(ws, "L38", str(datos.get("segundo_apellido", "")))
    escribir_celda(ws, "S38", str(datos.get("primer_nombre", "")))
    escribir_celda(ws, "Z38", str(datos.get("segundo_nombre", "")))
    escribir_celda(ws, "C43", str(datos.get("cedula", "")))
    escribir_celda(ws, "K43", str(datos.get("expedida_en", "")))
    escribir_celda(ws, "S43", str(datos.get("fecha_expedicion", "")))
    escribir_celda(ws, "AA43", str(datos.get("nacionalidad", "")))
    escribir_celda(ws, "I45", str(datos.get("direccion", "")))
    escribir_celda(ws, "U45", str(datos.get("barrio", "")))
    escribir_celda(ws, "H47", str(datos.get("departamento", "")))
    escribir_celda(ws, "O47", str(datos.get("municipio", "")))
    escribir_celda(ws, "Y47", str(datos.get("vereda", "")))
    escribir_celda(ws, "AA47", str(datos.get("correo", "")))
    escribir_celda(ws, "T49", str(datos.get("telefono", "")))
    escribir_celda(ws, "AA49", str(datos.get("celular", "")))
    escribir_celda(ws, "F52", str(datos.get("fecha_nacimiento", "")))
    escribir_celda(ws, "L51", calcular_edad(datos.get("fecha_nacimiento", "")))
    escribir_celda(ws, "AA55", str(datos.get("nivel_educativo", "")))
    escribir_celda(ws, "AC61", str(datos.get("n_hijos", "")))
    escribir_celda(ws, "H68", str(datos.get("cargo", "")))
    escribir_celda(ws, "S68", str(datos.get("empresa", "")))
    escribir_celda(ws, "H72", str(datos.get("ingreso_mensual", "")))
    escribir_celda(ws, "C74", str(datos.get("estrato", "")))
    escribir_celda(ws, "H88", str(datos.get("conducta_punible", "")))
    escribir_celda(ws, "I90", str(datos.get("numero_proceso", "")))
    escribir_celda(ws, "AA88", str(datos.get("fecha_hora_captura", "")))
    escribir_celda(ws, "H92", str(datos.get("fiscal", "")))
    escribir_celda(ws, "L92", str(datos.get("juez", "")))
    escribir_celda(ws, "A107", str(datos.get("resumen_hechos", "")))

    nombre_completo = f"{datos.get('primer_nombre','')} {datos.get('segundo_nombre','')} {datos.get('primer_apellido','')} {datos.get('segundo_apellido','')}".strip()
    escribir_celda(ws, "R116", "NO FIRMA PORQUE SE HIZO VIRTUAL")
    escribir_celda(ws, "Q119", str(datos.get("cedula", "")))
    escribir_celda(ws, "P119", str(datos.get("tipo_documento", "CC")).upper())

    tipo_doc = datos.get("tipo_documento", "").strip().lower()
    escribir_celda(ws, "F41", "X" if tipo_doc == "cc" else "")
    escribir_celda(ws, "I41", "X" if tipo_doc == "ti" else "")
    escribir_celda(ws, "K41", "X" if tipo_doc == "ce" else "")

    estado = datos.get("estado_civil", "").strip().lower()
    escribir_celda(ws, "R51", "X" if estado == "casado" else "")
    escribir_celda(ws, "T51", "X" if estado == "soltero" else "")
    escribir_celda(ws, "W51", "X" if estado == "viudo" else "")
    escribir_celda(ws, "Z51", "X" if estado == "separado" else "")
    escribir_celda(ws, "AC51", "X" if estado in ("union libre", "unión libre", "unionlibre") else "")
    escribir_celda(ws, "N61", str(datos.get("nombre_conyuge", "")) if estado in ("casado", "union libre", "unión libre", "unionlibre") else "")

    sexo = datos.get("sexo", "").strip().lower()
    escribir_celda(ws, "C53", "X" if sexo == "femenino" else "")
    escribir_celda(ws, "F53", "X" if sexo in ("masculino", "m") else "")

    laboral = datos.get("situacion_laboral", "").strip().lower()
    escribir_celda(ws, "L65", "X" if laboral in ("dependiente", "trabajador dependiente", "empleado") else "")
    escribir_celda(ws, "S65", "X" if laboral in ("independiente", "trabajador independiente", "indep") else "")
    escribir_celda(ws, "Y64", "X" if laboral in ("desempleado", "sin empleo") else "")
    escribir_celda(ws, "AC65", "X" if laboral in ("estudiante",) else "")

    vivienda = datos.get("vivienda", "").strip().lower()
    escribir_celda(ws, "I76", "X" if vivienda in ("propia", "propiedad", "propio") else "")
    escribir_celda(ws, "R76", "X" if vivienda in ("arrendada", "arrendar", "arriendo", "arrendado") else "")
    escribir_celda(ws, "M76", "X" if vivienda in ("familiar", "familiar/otros", "familiar ") else "")

    wb.save(salida)

def llenar_excel2(plantilla, salida, datos):
    wb = load_workbook(plantilla)
    ws = wb.active

    ws["D7"] = str(datos.get("fecha_recepcion", ""))
    ws["E9"] = str(datos.get("conducta_punible", ""))
    nombre_completo = f"{datos.get('primer_nombre','')} {datos.get('segundo_nombre','')} {datos.get('primer_apellido','')} {datos.get('segundo_apellido','')}".strip()
    ws["D8"] = nombre_completo
    ws["D15"] = str(datos.get("centro_reclusion", ""))

    privado = str(datos.get("privado_libertad", "false")).lower() in ("true", "1", "si", "sí")
    ws["F14"] = "X" if privado else ""
    ws["H14"] = "X" if not privado else ""

    ws["G40"] = "NO FIRMA PORQUE SE HIZO VIRTUAL"

    wb.save(salida)

def llenar_word(plantilla, salida, datos):
    doc = DocxTemplate(plantilla)
    context = {
        "primer_nombre": datos.get("primer_nombre", ""),
        "segundo_nombre": datos.get("segundo_nombre", ""),
        "primer_apellido": datos.get("primer_apellido", ""),
        "segundo_apellido": datos.get("segundo_apellido", ""),
        "nombre_completo": f"{datos.get('primer_nombre','')} {datos.get('segundo_nombre','')} {datos.get('primer_apellido','')} {datos.get('segundo_apellido','')}".strip(),
        "cedula": datos.get("cedula", ""),
        "tipo_documento": datos.get("tipo_documento", ""),
        "fecha_recepcion": datos.get("fecha_recepcion", ""),
        "lugar_recepcion": datos.get("lugar_recepcion", ""),
        "conducta_punible": datos.get("conducta_punible", ""),
        "centro_reclusion": datos.get("centro_reclusion", ""),
        "resumen_hechos": datos.get("resumen_hechos", ""),
        "firma": "NO FIRMA PORQUE SE HIZO VIRTUAL"
    }
    doc.render(context)
    doc.save(salida)
     
def convertir_docx_a_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        if DOCX2PDF_AVAILABLE:
            docx2pdf_convert(docx_path)
            if os.path.exists(pdf_path):
                return pdf_path
        else:
            # Fallback para Linux usando LibreOffice
            out_dir = os.path.dirname(os.path.abspath(docx_path))
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", os.path.abspath(docx_path), "--outdir", out_dir],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            if os.path.exists(pdf_path):
                return pdf_path
    except Exception as e:
        print("No se pudo convertir DOCX a PDF:", e)
    return None

def convertir_xlsx_a_pdf_windows(xlsx_path, output_pdf_path):
    if not WIN32_AVAILABLE:
        return False, "win32com no disponible"
    excel = None
    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = True
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(os.path.abspath(xlsx_path), UpdateLinks=False, ReadOnly=True)
        for sheet in wb.Worksheets:
            try:
                sheet.Visible = True
                sheet.Activate()
                excel.ActiveWindow.View = 1
                sheet.PageSetup.Zoom = False
                sheet.PageSetup.FitToPagesWide = 1
                sheet.PageSetup.FitToPagesTall = False
                time.sleep(0.2)
            except Exception:
                pass
        wb.ExportAsFixedFormat(
            Type=0,
            Filename=os.path.abspath(output_pdf_path),
            Quality=0,
            IncludeDocProperties=True,
            IgnorePrintAreas=False,
            OpenAfterPublish=False
        )
        wb.Close(SaveChanges=False)
        excel.Quit()
        return True, None
    except Exception as e:
        try:
            if excel: excel.Quit()
        except: pass
        return False, f"Error al exportar Excel a PDF: {e}"

def unir_pdfs(lista_pdfs, salida_final):
    merger = PdfMerger()
    for pdf in lista_pdfs:
        if pdf and os.path.exists(pdf):
            merger.append(pdf)
    merger.write(salida_final)
    merger.close()

def convertir_documento_api(file_path: str):
    api_secret = os.getenv("CONVERT_API_SECRET")
    if not api_secret:
        print("ADVERTENCIA: CONVERT_API_SECRET no está configurado.")
        return None
        
    ext = file_path.split('.')[-1].lower()
    if ext not in ['docx', 'xlsx']:
        return None
        
    url = f"https://v2.convertapi.com/convert/{ext}/to/pdf?Secret={api_secret}"
    try:
        with open(file_path, "rb") as f:
            files = {"File": f}
            response = requests.post(url, files=files, timeout=60)
            
        if response.status_code == 200:
            data = response.json()
            if "Files" in data and len(data["Files"]) > 0:
                file_data = base64.b64decode(data["Files"][0]["FileData"])
                out_path = file_path.replace(f".{ext}", ".pdf")
                with open(out_path, "wb") as out:
                    out.write(file_data)
                return out_path
        print(f"Error de ConvertAPI: {response.text}")
    except Exception as e:
        print(f"Error llamando a ConvertAPI: {e}")
    return None

@app.post("/api/generate")
async def generate_documents(datos: Dict[str, Any]):
    uid = str(uuid.uuid4())[:8]
    nombre_base = f"{datos.get('primer_nombre','sin_nombre')}_{datos.get('cedula','')}_{uid}".replace(" ", "_")
    
    excel1_out = os.path.join(RUTA_SALIDA, f"formato1_{nombre_base}.xlsx")
    excel2_out = os.path.join(RUTA_SALIDA, f"formato2_{nombre_base}.xlsx")
    word_out = os.path.join(RUTA_SALIDA, f"formato3_{nombre_base}.docx")
    
    formato1_in = os.path.join(RUTA_FORMATOS, "formato1.xlsx")
    formato2_in = os.path.join(RUTA_FORMATOS, "formato2.xlsx")
    formato3_in = os.path.join(RUTA_FORMATOS, "formato3.docx")
    
    if not os.path.exists(formato1_in) or not os.path.exists(formato2_in) or not os.path.exists(formato3_in):
        raise HTTPException(status_code=500, detail="Faltan las plantillas en la carpeta backend/formatos")

    try:
        # Fill templates
        llenar_excel1(formato1_in, excel1_out, datos)
        llenar_excel2(formato2_in, excel2_out, datos)
        llenar_word(formato3_in, word_out, datos)

        # Archivos generados correctamente (en /tmp)
        # Convertirlos a PDF vía ConvertAPI
        pdf_excel1 = convertir_documento_api(excel1_out)
        pdf_word = convertir_documento_api(word_out)
        pdf_excel2 = convertir_documento_api(excel2_out)
        
        # Unir PDFs en el orden solicitado: Formato 1, Formato 3, Formato 2
        orden = [pdf_excel1, pdf_word, pdf_excel2]
        orden_existentes = [p for p in orden if p and os.path.exists(p)]
        
        nombre_completo = f"{datos.get('primer_nombre','')} {datos.get('segundo_nombre','')} {datos.get('primer_apellido','')} {datos.get('segundo_apellido','')}".strip().replace(" ", "_").upper()
        if not nombre_completo: nombre_completo = "SIN_NOMBRE"
        
        salida_final_name = f"FICHA_{nombre_completo}.pdf"
        salida_final = os.path.join(RUTA_SALIDA, salida_final_name)
        
        if len(orden_existentes) > 0:
            unir_pdfs(orden_existentes, salida_final)
        else:
            raise HTTPException(status_code=500, detail="No se pudo convertir ningún documento a PDF. Verifica tu CONVERT_API_SECRET.")
            
        if not os.path.exists(salida_final):
            raise HTTPException(status_code=500, detail="No se pudo crear el archivo PDF consolidado")

        # Devolver el archivo directamente
        return FileResponse(salida_final, media_type='application/pdf', filename=salida_final_name)

    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=f"Error durante la generación: {str(e)}\n{traceback.format_exc()}")

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(RUTA_SALIDA, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    return FileResponse(file_path, media_type='application/pdf', filename=filename)

@app.get("/api/template")
async def download_template():
    template_path = os.path.join(RUTA_FORMATOS, "PLANTILLA_INGRESO_DATOS.docx")
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    return FileResponse(
        template_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename="Plantilla_Defensoria.docx"
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
